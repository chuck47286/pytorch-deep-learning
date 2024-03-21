import re
from docx import Document
import os


def print_document_headings(doc_path):
    """
    打印指定Word文档中的所有标题。

    参数:
    - doc_path: Word文档的路径。
    """
    # 打开Word文档
    doc = Document(doc_path)

    # 遍历文档中的每个段落
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        style_name = paragraph.style.name
        # print(f"Text: {text[:30]}... Style: {style_name}")
        # 检查段落的样式是否为标题样式
        if paragraph.style.name.startswith('Heading'):
            print(paragraph.text)


def print_all_tables(doc_path):
    """
    打印Word文档中所有表格的内容。

    参数:
    - doc_path: Word文档的路径。
    """
    doc = Document(doc_path)

    # 遍历文档中的所有表格
    for table_index, table in enumerate(doc.tables, start=1):
        # 过滤表格，方式1（直接指定表格下标过滤）
        # 如果表格索引不在指定范围内，则跳过该表格(设置开关,原则：前开后闭)
        # start_table_index, end_table_index = 35, 44
        # if table_index < start_table_index or table_index >= end_table_index:
        #     continue

        # 过滤表格，方式2 （通过table表格头的样式内容校验）
        if len(table.rows) > 0:
            header_cells = table.rows[0].cells
            header_contents = [cell.text.strip() for cell in header_cells]
            # 检查表头是否符合预期
            # 定义所有可能的表头
            possible_headers = [
                ["标签", "字段名", "字段描述", "必须", "类型"],
                ["标签", "字段名", "字段描述", "必填", "类型"]
            ]
            # 检查当前表格的表头是否符合预期的任何一种
            header_matches = any(
                all(header in header_contents for header in expected_headers)
                for expected_headers in possible_headers
            )
            if not header_matches:
                continue

        print(f"Table {table_index}:")
        # 遍历表格中的每一行
        for row in table.rows:
            row_data = []
            previous_cell_text = None  # 初始化前一个单元格的内容变量

            # 遍历行中的每一个单元格，并收集单元格的文本
            for cell in row.cells:
                # row_data.append(cell.text.strip()) # 之前逻辑，没有单元格内容校验

                cell_text = cell.text.strip()
                # 检查当前单元格内容是否与前一个单元格的内容一致
                if cell_text != previous_cell_text:
                    row_data.append(cell_text)

                # 更新前一个单元格的内容
                previous_cell_text = cell_text

            # 将单元格文本连接成一行，并打印
            print(" | ".join(row_data))

        # 在表格之间添加空行以提高可读性
        print("\n" + "-" * 50 + "\n")


def remove_consecutive_duplicates(items):
    """
    移除列表中连续重复的元素。
    参数:
    - items: 要处理的元素列表。

    返回:
    - 新列表，其中连续重复的元素被移除。
    :param items:
    :return:
    """
    if not items:
        return []

    # 初始化结果列表并设置第一个元素
    result = [items[0]]

    # 遍历剩余元素，只添加与前一个不同的元素
    for item in items[1:]:
        if item != result[-1]:  # 与结果列表中最后一个元素比较
            result.append(item)

    return result


# def remove_elements_if_contains_special_character(row_data):
#     """
#     如果list中存在不合理的字段（'→'），则需要删除该字段
#     :param row_data: 原始行数据列表
#     :return: 调整后的行数据列表
#     """
#     # 使用正则表达式来检查是否为数字
#     is_digit = re.compile(r'^\d+$')
#
#     # 确保列表至少有两个元素
#     if len(row_data) >= 2:
#         first_is_digit = is_digit.match(row_data[0])
#         second_is_digit = is_digit.match(row_data[1])
#
#         # 如果第一个元素是数字且第二个元素不是数字，则交换这两个元素
#         if first_is_digit and not second_is_digit:
#             row_data[0], row_data[1] = row_data[1], row_data[0]
#
#     return row_data


def clean_row_data(row_data):
    """
    清洗数据
    按照文档的格式以及内容来，规则化数据
    :param row_data:
    :return:
    """
    original_row_data = row_data
    # 0.规则化循环体数据，直接删除→字段
    row_data = [item for item in row_data if item != '→']
    # 1.如果row_data数据小于5，可以直接删除
    if len(row_data) < 5:
        print(f"row_data < 5，可以直接删除.({original_row_data})")
        return []
    # 2.如果row_data数据大于5，需要删除头部非数字数据
    tag_pattern = re.compile(r'^\d+$')
    if len(row_data) > 5 and not tag_pattern.match(row_data[0]):
        row_data = row_data[1:]

    if len(row_data) != 5:
        print(f"row_data > 5，需要删除头部非数字数据后，数据不满足长度要求，所以要删除.({original_row_data})")
        return []
    # 3.如果row_data数据==5，可是首项和尾项的元素为空，也要清除
    if not row_data[0] or not row_data[-1]:
        print(f"row_data == 5，可是首项和尾项的元素为空，也要清除. ({original_row_data})")
        return []
    # 通过数据清洗，直接返回
    return row_data


def tables_to_structured_data_simplified(doc_path):
    """
    将表格数据格式化
    :param doc_path:
    :return:
    """
    doc = Document(doc_path)
    structured_data = []  # 用于保存所有表格数据的列表

    # 遍历文档中的所有表格
    for table_index, table in enumerate(doc.tables, start=1):
        # 过滤表格，方式1（直接指定表格下标过滤）
        # 如果表格索引不在指定范围内，则跳过该表格(设置开关,原则：前开后闭)
        # start_table_index, end_table_index = 52, 54
        # if table_index < start_table_index or table_index >= end_table_index:
        #     continue

        # 过滤表格，方式2 （通过table表格头的样式内容校验）
        if len(table.rows) > 0:
            header_cells = table.rows[0].cells
            header_contents = [cell.text.strip() for cell in header_cells]
            # 检查表头是否符合预期
            # 定义所有可能的表头
            possible_headers = [
                ["标签", "字段名", "字段描述", "必须", "类型"],
                ["标签", "字段名", "字段描述", "必填", "类型"]
            ]
            # 检查当前表格的表头是否符合预期的任何一种
            header_matches = any(
                all(header in header_contents for header in expected_headers)
                for expected_headers in possible_headers
            )
            if not header_matches:
                continue

        table_data = {}  # 当前表格的数据将被保存在这个字典中
        table_name = None  # 用于保存当前表格的名称（从第二行的某列获取）
        # 遍历表格中的每一行
        for row_index, row in enumerate(table.rows):
            if row_index == 0:
                continue  # 跳过表头

            # 数据去重，主要是由于文档合并文档操作所需的
            row_data = remove_consecutive_duplicates([cell.text.strip() for cell in row.cells])
            if row_index == 1:
                # 使用next()查找报告“MsgType”的单元格，否则使用默认名称
                table_name = next((cell.split("=")[1].strip() for cell in row_data if "MsgType" and "=" in cell),
                                  "DefaultTableName")
                table_data[table_name] = []
                print(f"Table {table_index}: TableName={table_name}")  # 标识出当前处理的table下标

            if row_index > 1:  # 从第三行开始保存数据
                # 数据清洗（不满足5列数据的需要规则化）
                row_data = clean_row_data(row_data)
                if not row_data:
                    continue
                key = row_data[0]  # 第一列作为键
                values = row_data[1:]  # 剩余作为值
                table_data[table_name].append({key: values})

        structured_data.append(table_data)

    return structured_data


def check_structured_data_and_find_differences(structured_data):
    """
    用于统计消息类型中的tag的类型是否都是一致的
    :param structured_data:
    :return:
    """
    # 使用字典收集每个tag的最后一列值，同时记录消息类型
    tag_values_dict = {}
    list1 = []
    list2 = []
    list3 = []

    for table_data in structured_data:
        for msg_type, entries in table_data.items():
            for entry in entries:
                for tag, values in entry.items():
                    if tag not in tag_values_dict:
                        tag_values_dict[tag] = {}
                    tag_values_dict[tag].setdefault(msg_type, set()).add(values[-1])

    for tag, msg_types in tag_values_dict.items():
        combined_values = {v for values in msg_types.values() for v in values}
        msg_types_join = ','.join(msg_types.keys())
        if len(combined_values) > 1:
            # print(f"Tag '{tag}' 在不同消息类型中有不一致的最后一列值: {msg_types}")
            list1.append(f"Tag '{tag}' 在不同消息类型中有不一致的最后一列值: {msg_types}")

        elif len(msg_types) > 1:
            # print(f"Tag '{tag}' 出现在多个消息类型({msg_types_join})中，但其值一致。")
            list2.append(f"Tag '{tag}' 出现在多个消息类型({msg_types_join})中，但其值一致。")
        else:
            # 如果只有一个消息类型包含这个tag,且想展示这种情况
            # print(f"Tag '{tag}' 仅出现在消息类型({msg_types_join})")
            list3.append(f"Tag '{tag}' 仅出现在消息类型({msg_types_join})")

    [print(item) for item in list1]
    print(len(list1))
    [print(item) for item in list2]
    print(len(list2))
    [print(item) for item in list3]
    print(len(list3))
    print(
        f"---------------\n"
        f"Total:{len(list1) + len(list2) + len(list3)}, right={len(list2) + len(list3)}, mistake={len(list1)}"
        f"\n-------------")


if __name__ == "__main__":
    # 读取文件路径
    dir_path = r'../external_library/doc_folder/'  # 测试文件路径
    file_name = "20240314.docx"  # 测试文件名

    file_path = os.path.join(dir_path, file_name)
    # print_document_headings(file_path)

    # print_all_tables(file_path)  # 打印所有的表格内容

    data = tables_to_structured_data_simplified(file_path)  # 打印所有的格式化后的表格内容
    for table in data:
        print(table)
    print(len(data))
    print("--------------")
    check_structured_data_and_find_differences(data)
